import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

# ==================== HELPER FUNCTIONS ====================

def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_chapter_heading(doc, text):
    doc.add_page_break()
    heading = doc.add_heading(text, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading.paragraph_format.space_before = Pt(36)
    heading.paragraph_format.space_after = Pt(24)
    heading.paragraph_format.keep_with_next = True
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(30, 58, 138) # Navy Blue
    return heading

def add_section_heading(doc, text):
    heading = doc.add_heading(text, level=2)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading.paragraph_format.space_before = Pt(24)
    heading.paragraph_format.space_after = Pt(12)
    heading.paragraph_format.keep_with_next = True
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = RGBColor(234, 88, 12) # Saffron
    return heading

def add_subsection_heading(doc, text):
    heading = doc.add_heading(text, level=3)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading.paragraph_format.space_before = Pt(18)
    heading.paragraph_format.space_after = Pt(6)
    heading.paragraph_format.keep_with_next = True
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.italic = True
        run.font.color.rgb = RGBColor(31, 41, 55) # Dark Gray
    return heading

# Add Formatted Paragraph with Double Spacing (2.0)
def add_paragraph(doc, text="", bold_prefix=""):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0  # Academic Double Spacing
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Times New Roman'
        r_pre.font.size = Pt(12)
        r_pre.font.bold = True
        r_pre.font.color.rgb = RGBColor(31, 41, 55)
        
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(31, 41, 55)
    return p

# Add Bullet/Numbered Point
def add_bullet_point(doc, text, bold_prefix="", num_style=False):
    style_name = 'List Number' if num_style else 'List Bullet'
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.line_spacing = 1.6
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Times New Roman'
        r_pre.font.size = Pt(12)
        r_pre.font.bold = True
        
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(31, 41, 55)
    return p

# Add Diagram Image Centered
def add_diagram(doc, img_name, caption):
    if os.path.exists(img_name):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(18)
        p_img.paragraph_format.space_after = Pt(6)
        r_img = p_img.add_run()
        r_img.add_picture(img_name, width=Inches(6.2))
        
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_before = Pt(6)
        p_cap.paragraph_format.space_after = Pt(24)
        r_cap = p_cap.add_run(f"Figure: {caption}")
        r_cap.font.name = 'Times New Roman'
        r_cap.font.size = Pt(10)
        r_cap.font.italic = True
        r_cap.font.color.rgb = RGBColor(100, 110, 120)
    else:
        add_paragraph(doc, f"[DIAGRAM PLACEHOLDER: {img_name} - {caption}]")

# Create Styled Table
def create_styled_table(doc, rows, cols, headers, col_widths=None):
    table = doc.add_table(rows=rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Format Header Row
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "1E3A8A") # Navy Blue Header
        set_cell_margins(hdr_cells[i], top=150, bottom=150, left=150, right=150)
        
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255) # White Text
            
    # Apply column widths
    if col_widths:
        for r in table.rows:
            for i, w in enumerate(col_widths):
                r.cells[i].width = Inches(w)
                
    return table

# Add Row to Table
def add_table_row(table, data, is_align_center=False):
    row_cells = table.add_row().cells
    for i, val in enumerate(data):
        row_cells[i].text = str(val)
        set_cell_margins(row_cells[i], top=100, bottom=100, left=120, right=120)
        p = row_cells[i].paragraphs[0]
        p.paragraph_format.line_spacing = 1.35
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        
        if is_align_center or i == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(31, 41, 55)

def build_thesis_docx():
    print("Initializing expanded document (Target: 65+ pages with updated cover details)...")
    doc = Document()
    
    # Page Setup
    section = doc.sections[0]
    section.page_width = Inches(8.27)  # A4 size
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.3) # binding margin
    section.right_margin = Inches(1.0)
    
    # Configure Default Paragraph Format
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Times New Roman'
    font_normal.size = Pt(12)
    font_normal.color.rgb = RGBColor(31, 41, 55)
    style_normal.paragraph_format.line_spacing = 2.0
    style_normal.paragraph_format.space_before = Pt(6)
    style_normal.paragraph_format.space_after = Pt(18)
    
    # ==================== COVER PAGE 1 (LEFT SIDE SCREENSHOT) ====================
    print("Writing First Cover Page...")
    p_logo_space = doc.add_paragraph()
    p_logo_space.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo_space.paragraph_format.space_before = Pt(30)
    
    p_t1 = doc.add_paragraph()
    p_t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t1 = p_t1.add_run("ONLINE-EXAMINATION-SYSTEM")
    r_t1.font.name = 'Times New Roman'
    r_t1.font.size = Pt(22)
    r_t1.font.bold = True
    r_t1.font.color.rgb = RGBColor(234, 88, 12) # Saffron/Red tone
    p_t1.paragraph_format.space_after = Pt(40)
    
    # DEVELOPED BY details
    p_dev1 = doc.add_paragraph()
    p_dev1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_dev_lbl1 = p_dev1.add_run("DEVELOPED BY\n")
    r_dev_lbl1.font.bold = True
    r_dev_lbl1.font.size = Pt(11)
    r_dev_val1 = p_dev1.add_run("(MAYUR RAMAVAT) (SUK230704CA062)\n\n")
    r_dev_val1.font.bold = True
    r_dev_val1.font.size = Pt(12)
    
    r_guid_lbl1 = p_dev1.add_run("UNDER GUIDANCE OF\n")
    r_guid_lbl1.font.bold = True
    r_guid_lbl1.font.size = Pt(11)
    r_guid_val1 = p_dev1.add_run("MR. DIPESH DAVE\nASSISTANT PROFESSOR\nFACULTY OF IT AND COMPUTER SCIENCE\n\n")
    r_guid_val1.font.bold = True
    r_guid_val1.font.size = Pt(12)
    
    r_sub_to_lbl1 = p_dev1.add_run("SUBMITTED TO\n")
    r_sub_to_lbl1.font.bold = True
    r_sub_to_lbl1.font.size = Pt(11)
    p_dev1.paragraph_format.space_after = Pt(30)
    
    p_univ1 = doc.add_paragraph()
    p_univ1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_univ_text1 = p_univ1.add_run("FACULTY OF IT AND COMPUTER SCIENCE\nSWAMINARAYAN UNIVERSITY\nSAIJ, KALOL\n\n")
    r_univ_text1.font.bold = True
    r_univ_text1.font.size = Pt(13)
    r_univ_text1.font.color.rgb = RGBColor(30, 58, 138)
    
    r_fullfill1 = p_univ1.add_run("FOR PARTIAL FULFILLMENT TOWARDS THE AWARD OF\nIN\n")
    r_fullfill1.font.size = Pt(11)
    r_degree1 = p_univ1.add_run("BACHELOR OF COMPUTER APPLICATIONS- (B.C.A)\n")
    r_degree1.font.bold = True
    r_degree1.font.size = Pt(12)
    r_degree1.font.color.rgb = RGBColor(234, 88, 12)
    
    r_date1 = p_univ1.add_run("JUNE - 2026")
    r_date1.font.bold = True
    r_date1.font.size = Pt(11)
    
    # ==================== COVER PAGE 2 (RIGHT SIDE SCREENSHOT) ====================
    print("Writing Second Cover Page...")
    doc.add_page_break()
    p_logo_space2 = doc.add_paragraph()
    p_logo_space2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_logo_space2.paragraph_format.space_before = Pt(30)
    
    p_t2 = doc.add_paragraph()
    p_t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t2_lbl = p_t2.add_run("A\nPROJECT REPORT\nON\n\n")
    r_t2_lbl.font.size = Pt(12)
    r_t2_lbl.font.bold = True
    r_t2_lbl.font.color.rgb = RGBColor(234, 88, 12)
    r_t2 = p_t2.add_run("ONLINE-EXAMINATION-SYSTEM")
    r_t2.font.name = 'Times New Roman'
    r_t2.font.size = Pt(22)
    r_t2.font.bold = True
    r_t2.font.color.rgb = RGBColor(30, 58, 138)
    p_t2.paragraph_format.space_after = Pt(40)
    
    # Developed by / Guide parallel text layout using table for cover formatting
    tbl_cover = doc.add_table(rows=1, cols=2)
    tbl_cover.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_cover.style = 'Normal Table'
    
    c_dev = tbl_cover.rows[0].cells[0]
    p_c_dev = c_dev.paragraphs[0]
    p_c_dev.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_c_dev_lbl = p_c_dev.add_run("DEVELOPED BY\n")
    r_c_dev_lbl.font.bold = True
    r_c_dev_lbl.font.size = Pt(11)
    r_c_dev_lbl.font.color.rgb = RGBColor(234, 88, 12)
    r_c_dev_val = p_c_dev.add_run("(MAYUR RAMAVAT) (SUK230704CA062)")
    r_c_dev_val.font.bold = True
    r_c_dev_val.font.size = Pt(11)
    
    c_guid = tbl_cover.rows[0].cells[1]
    p_c_guid = c_guid.paragraphs[0]
    p_c_guid.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_c_guid_lbl = p_c_guid.add_run("INTERNAL GUIDE\n")
    r_c_guid_lbl.font.bold = True
    r_c_guid_lbl.font.size = Pt(11)
    r_c_guid_lbl.font.color.rgb = RGBColor(234, 88, 12)
    r_c_guid_val = p_c_guid.add_run("MR. DIPESH DAVE")
    r_c_guid_val.font.bold = True
    r_c_guid_val.font.size = Pt(11)
    
    p_univ2 = doc.add_paragraph()
    p_univ2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_univ2.paragraph_format.space_before = Pt(40)
    r_univ_text2 = p_univ2.add_run("\nFACULTY OF IT AND COMPUTER SCIENCE\nSWAMINARAYAN UNIVERSITY\n\n")
    r_univ_text2.font.bold = True
    r_univ_text2.font.size = Pt(14)
    r_univ_text2.font.color.rgb = RGBColor(30, 58, 138)
    
    r_thesis2 = p_univ2.add_run("A Project Thesis Submitted to\nSwaminarayan University\nIn Partial Fulfillment towards the Award of Degree\nIn ")
    r_thesis2.font.size = Pt(11)
    r_degree2 = p_univ2.add_run("BACHELOR OF COMPUTER APPLICATIONS- (B.C.A)\n")
    r_degree2.font.bold = True
    r_degree2.font.size = Pt(12)
    r_degree2.font.color.rgb = RGBColor(234, 88, 12)
    
    r_date2 = p_univ2.add_run("JUNE - 2026\n\n")
    r_date2.font.bold = True
    r_date2.font.size = Pt(11)
    
    p_univ2_foot = doc.add_paragraph()
    p_univ2_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_foot2 = p_univ2_foot.add_run("FACULTY OF IT AND COMPUTER SCIENCE\nSWAMINARAYAN UNIVERSITY\nSAIJ, KALOL")
    r_foot2.font.bold = True
    r_foot2.font.size = Pt(12)
    r_foot2.font.color.rgb = RGBColor(30, 58, 138)
    
    # ==================== CERTIFICATE 1 (INTERNAL GUIDE - LEFT SIDE) ====================
    print("Writing Certificate Page 1...")
    doc.add_page_break()
    p_cert1_hdr = doc.add_paragraph()
    p_cert1_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cert1_hdr.paragraph_format.space_before = Pt(20)
    r_c1_hdr = p_cert1_hdr.add_run("SWAMINARAYAN UNIVERSITY\nFACULTY OF IT AND COMPUTER SCIENCE\n\n")
    r_c1_hdr.font.bold = True
    r_c1_hdr.font.size = Pt(16)
    r_c1_hdr.font.color.rgb = RGBColor(234, 88, 12)
    
    r_c1_lbl = p_cert1_hdr.add_run("CERTIFICATE")
    r_c1_lbl.font.bold = True
    r_c1_lbl.font.size = Pt(18)
    r_c1_lbl.font.color.rgb = RGBColor(30, 58, 138)
    p_cert1_hdr.paragraph_format.space_after = Pt(30)
    
    add_paragraph(doc, "This is to certify that the project report entitled \"ONLINE-EXAMINATION-SYSTEM\" submitted in partial fulfillment of the requirements for the award of the degree of B.C.A in FACULTY OF IT AND COMPUTER SCIENCE, SWAMINARAYAN UNIVERSITY(SU), SAIJ, KALOL is a Bonafide work carried out by MAYUR RAMAVAT, Enrollment No SUK230704CA062, during the academic year 2025-2026.")
    
    p_sig1_space = doc.add_paragraph()
    p_sig1_space.paragraph_format.space_before = Pt(120)
    p_sig1_space.paragraph_format.line_spacing = 1.3
    p_sig1_space.add_run("DATE:\n\n")
    r_sig1_f = p_sig1_space.add_run("SIGN OF INTERNAL GUIDE:\t\t\t\t\tSeal of Institute")
    r_sig1_f.font.bold = True
    
    # ==================== CERTIFICATE 2 (EXTERNAL EXAMINER - RIGHT SIDE) ====================
    print("Writing Certificate Page 2...")
    doc.add_page_break()
    p_cert2_hdr = doc.add_paragraph()
    p_cert2_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cert2_hdr.paragraph_format.space_before = Pt(20)
    r_c2_hdr = p_cert2_hdr.add_run("SWAMINARAYAN UNIVERSITY\nFACULTY OF IT AND COMPUTER SCIENCE\n\n")
    r_c2_hdr.font.bold = True
    r_c2_hdr.font.size = Pt(16)
    r_c2_hdr.font.color.rgb = RGBColor(234, 88, 12)
    
    r_c2_lbl = p_cert2_hdr.add_run("CERTIFICATE")
    r_c2_lbl.font.bold = True
    r_c2_lbl.font.size = Pt(18)
    r_c2_lbl.font.color.rgb = RGBColor(30, 58, 138)
    p_cert2_hdr.paragraph_format.space_after = Pt(30)
    
    add_paragraph(doc, "This is to certify that the project report entitled \"ONLINE-EXAMINATION-SYSTEM\" submitted in partial fulfillment of the requirements for the award of the degree of B.C.A in FACULTY OF IT AND COMPUTER SCIENCE, SWAMINARAYAN UNIVERSITY(SU), SAIJ, KALOL is a Bonafide work carried out by MAYUR RAMAVAT, Enrollment No SUK230704CA062, during the academic year 2025-2026.")
    
    p_sig2_space = doc.add_paragraph()
    p_sig2_space.paragraph_format.space_before = Pt(120)
    p_sig2_space.paragraph_format.line_spacing = 1.3
    p_sig2_space.add_run("DATE:\n\n")
    r_sig2_f = p_sig2_space.add_run("SIGN OF EXTERNAL:\t\t\t\t\t\tSeal of Institute")
    r_sig2_f.font.bold = True
    
    # ==================== ACKNOWLEDGEMENT ====================
    print("Writing Acknowledgement...")
    doc.add_page_break()
    p_ack_title = doc.add_paragraph()
    p_ack_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ack_title = p_ack_title.add_run("ACKNOWLEDGEMENT")
    r_ack_title.font.name = 'Times New Roman'
    r_ack_title.font.size = Pt(16)
    r_ack_title.font.bold = True
    r_ack_title.font.color.rgb = RGBColor(30, 58, 138)
    p_ack_title.paragraph_format.space_after = Pt(30)
    
    add_paragraph(doc, "I would like to express my deepest gratitude to my project guide, Mr. Dipesh Dave, Assistant Professor in the Faculty of IT and Computer Science, Swaminarayan University, for his invaluable guidance, continuous support, and technical feedback throughout the development of the Online Examination System. His insights into student assessment systems and auto-grading logic helped shape the project features.")
    add_paragraph(doc, "I extend my sincere thanks to the Department Heads of the Faculty of IT and Computer Science, for providing excellent academic facilities, computer laboratories, and server resources which enabled me to compile, run, and test the multi-tier application environment.")
    add_paragraph(doc, "I also express my appreciation to all my teachers, class mates, and family members for their motivation, technical discussions, and encouragement during the course of this project development. Their feedback and critiques were instrumental in perfecting the final build of the application.")
    
    p_ack_student = doc.add_paragraph()
    p_ack_student.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_ack_student.paragraph_format.space_before = Pt(80)
    r_ack_std1 = p_ack_student.add_run("Mayur Ramavat\n")
    r_ack_std1.font.bold = True
    r_ack_std2 = p_ack_student.add_run("Enrollment No: SUK230704CA062\nB.C.A Semester VI")
    
    # ==================== INDEX (TABLE OF CONTENTS) ====================
    print("Writing Index...")
    doc.add_page_break()
    p_idx_title = doc.add_paragraph()
    p_idx_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_idx_title = p_idx_title.add_run("INDEX / TABLE OF CONTENTS")
    r_idx_title.font.name = 'Times New Roman'
    r_idx_title.font.size = Pt(16)
    r_idx_title.font.bold = True
    r_idx_title.font.color.rgb = RGBColor(30, 58, 138)
    p_idx_title.paragraph_format.space_after = Pt(18)
    
    tbl_idx = create_styled_table(doc, 1, 3, ["SR. NO.", "CONTENT", "PAGE NO."], [1.0, 5.0, 1.0])
    
    index_items = [
        ("I", "Cover Page", ""),
        ("II", "First Page", ""),
        ("III", "Certificate", ""),
        ("", "Completion Certificate", ""),
        ("IV", "Acknowledgement", ""),
        ("V", "Index", ""),
        ("VI", "List of Figures", ""),
        ("VII", "List of Tables", ""),
        ("VIII", "List of Abbreviation", ""),
        ("IX", "Abstract", ""),
        ("1", "Introduction", ""),
        ("", "  1.1 Project Profile", ""),
        ("", "  1.2 Project Summary", ""),
        ("", "  1.3 purpose: Goals & Objectives", ""),
        ("", "  1.4 General Description", ""),
        ("2", "System Analysis", ""),
        ("", "  2.1 Study of Current System", ""),
        ("", "  2.2 Problems of Current System", ""),
        ("", "  2.3 Requirements of New System", ""),
        ("", "  2.4 Feasibility Study", ""),
        ("3", "System Requirement Study", ""),
        ("", "  3.1 User Characteristics", ""),
        ("", "    3.1.1 Visitor", ""),
        ("", "    3.1.2 User", ""),
        ("", "    3.1.3 Admin", ""),
        ("", "  3.2 Software & Hardware Requirement", ""),
        ("", "    3.2.1 Minimum Hardware Requirement", ""),
        ("", "    3.2.2 Minimum Software Requirement", ""),
        ("", "  3.3 Functional and Non functional requirement", ""),
        ("4", "System Design", ""),
        ("", "  4.1 System Diagram", ""),
        ("", "  4.2 Activity Diagram", ""),
        ("", "  4.3 Data Flow Diagram", ""),
        ("", "  4.4 Use Case Diagram", ""),
        ("", "  4.5 Data Modelling", ""),
        ("", "    4.5.1 E-R Diagram", ""),
        ("", "    4.5.2 Data Dictionary", ""),
        ("5", "Project Planning & Scheduling", ""),
        ("", "  5.1 Process model", ""),
        ("", "  5.2 Project Plan", ""),
        ("", "  5.3 Schedule Representation", ""),
        ("6", "Risk Management", ""),
        ("", "  6.1 Risk Identification", ""),
        ("", "  6.2 Risk Analysis", ""),
        ("", "  6.3 Risk Planning", ""),
        ("7", "Security Features", ""),
        ("8", "Testing", ""),
        ("", "  8.1 Testing Strategy", ""),
        ("", "  8.2 Testing Methods", ""),
        ("9", "Screen Shots", ""),
        ("", "  9.1 Admin Panel Screen Shots", ""),
        ("", "  9.2 User Panel Screen Shots", ""),
        ("10", "Conclusion", ""),
        ("11", "Future enhancement", ""),
        ("12", "Bibliography", "")
    ]
    
    for sr, item, pg in index_items:
        add_table_row(tbl_idx, [sr, item, pg])
        
    # ==================== LIST OF FIGURES & TABLES ====================
    print("Writing Lists of Figures and Tables...")
    doc.add_page_break()
    p_lof_title = doc.add_paragraph()
    r_lof_title = p_lof_title.add_run("LIST OF FIGURES")
    r_lof_title.font.bold = True
    r_lof_title.font.color.rgb = RGBColor(30, 58, 138)
    p_lof_title.paragraph_format.space_after = Pt(8)
    
    tbl_lof = create_styled_table(doc, 1, 3, ["FIGURE NO.", "FIGURE CAPTION", "PAGE NO."], [1.2, 4.8, 1.0])
    add_table_row(tbl_lof, ["Figure 4.1", "System Architecture (Three-Tier Diagram)", ""])
    add_table_row(tbl_lof, ["Figure 4.2", "Student Assessment & Proctoring Activity Flow", ""])
    add_table_row(tbl_lof, ["Figure 4.3", "Data Flow Diagram (DFD Level 1)", ""])
    add_table_row(tbl_lof, ["Figure 4.4", "Use Case Interaction Model", ""])
    add_table_row(tbl_lof, ["Figure 4.5", "Entity-Relationship (E-R) Database Schema", ""])
    
    p_lot_title = doc.add_paragraph()
    p_lot_title.paragraph_format.space_before = Pt(30)
    r_lot_title = p_lot_title.add_run("LIST OF TABLES")
    r_lot_title.font.bold = True
    r_lot_title.font.color.rgb = RGBColor(30, 58, 138)
    p_lot_title.paragraph_format.space_after = Pt(8)
    
    tbl_lot = create_styled_table(doc, 1, 3, ["TABLE NO.", "TABLE DESCRIPTION", "PAGE NO."], [1.2, 4.8, 1.0])
    add_table_row(tbl_lot, ["Table 1.1", "Project Profile Parameters", ""])
    add_table_row(tbl_lot, ["Table 3.1", "Minimum Hardware Configurations", ""])
    add_table_row(tbl_lot, ["Table 3.2", "Minimum Software Toolsets", ""])
    add_table_row(tbl_lot, ["Table 4.1", "Data Dictionary - Admins Table", ""])
    add_table_row(tbl_lot, ["Table 4.2", "Data Dictionary - Students Table", ""])
    add_table_row(tbl_lot, ["Table 4.3", "Data Dictionary - Exams Table", ""])
    add_table_row(tbl_lot, ["Table 4.4", "Data Dictionary - Questions Table", ""])
    add_table_row(tbl_lot, ["Table 4.5", "Data Dictionary - Student Exams Table", ""])
    add_table_row(tbl_lot, ["Table 4.6", "Data Dictionary - Student Answers Table", ""])
    add_table_row(tbl_lot, ["Table 6.1", "Risk Probability and Impact Matrix", ""])
    add_table_row(tbl_lot, ["Table 8.1", "Functional Test Cases and Verification Results", ""])
    
    # ==================== LIST OF ABBREVIATIONS ====================
    print("Writing Abbreviations...")
    doc.add_page_break()
    p_ab_title = doc.add_paragraph()
    r_ab_title = p_ab_title.add_run("LIST OF ABBREVIATION")
    r_ab_title.font.name = 'Times New Roman'
    r_ab_title.font.size = Pt(16)
    r_ab_title.font.bold = True
    r_ab_title.font.color.rgb = RGBColor(30, 58, 138)
    p_ab_title.paragraph_format.space_after = Pt(12)
    
    tbl_ab = create_styled_table(doc, 1, 2, ["ABBREVIATION", "EXPANDED FORM"], [2.0, 5.0])
    abbrevs = [
        ("MCQ", "Multiple Choice Question"),
        ("SQL", "Structured Query Language"),
        ("PHP", "Hypertext Preprocessor"),
        ("NLP", "Natural Language Processing"),
        ("DFD", "Data Flow Diagram"),
        ("ERD", "Entity Relationship Diagram"),
        ("SDLC", "Software Development Life Cycle"),
        ("CSS", "Cascading Style Sheets"),
        ("HTML", "Hypertext Markup Language"),
        ("JS", "JavaScript"),
        ("PDO", "PHP Data Objects"),
        ("API", "Application Programming Interface"),
        ("UI/UX", "User Interface / User Experience"),
        ("Gantt", "Gantt Chart Project Scheduler"),
        ("HOD", "Head of Department"),
        ("B.C.A", "Bachelor of Computer Applications")
    ]
    for ab, exp in abbrevs:
        add_table_row(tbl_ab, [ab, exp])
        
    # ==================== ABSTRACT ====================
    print("Writing Abstract...")
    doc.add_page_break()
    p_abs_title = doc.add_paragraph()
    p_abs_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_abs_title = p_abs_title.add_run("ABSTRACT")
    r_abs_title.font.name = 'Times New Roman'
    r_abs_title.font.size = Pt(16)
    r_abs_title.font.bold = True
    r_abs_title.font.color.rgb = RGBColor(30, 58, 138)
    p_abs_title.paragraph_format.space_after = Pt(16)
    
    add_paragraph(doc, "Online assessment systems have become a cornerstone of modern educational infrastructure. However, standard examination platforms suffer from two critical limitations: first, they are largely restricted to multiple-choice questions (MCQs), failing to assess students' subjective critical thinking; second, they lack robust proctoring mechanisms to detect web-based search cheating and window switching. This project presents an Online Examination System with Semantic Evaluation and Real-Time Proctoring designed to bridge these gaps.")
    add_paragraph(doc, "The system is structured as a three-tier web application built on PHP, MySQL, and JavaScript, integrated with an offline Python-based NLP grading engine. For security and integrity, the client-side browser disables right-click, blocks copy-paste keyboard shortcut events (Ctrl+C, Ctrl+V, Ctrl-X, Ctrl-A), and binds event listeners to window blur and page visibility changes to detect tab-switching. If a student switches tabs or exits focus, the system logs the incident and triggers live supervisor warnings. On exceeding five violations, the active exam is immediately locked and force-submitted to the server.")
    add_paragraph(doc, "For subjective assessment, descriptive questions are graded automatically by a Python module using a hybrid scoring algorithm. The engine computes similarity indices by blending character-level sequence matching (SequenceMatcher) with key concept density overlap, scaling results to match instructor model answers without requiring verbatim matches. Admin interfaces display a live monitor widget that auto-refreshes every 5 seconds, allowing supervisors to oversee active students, review logged violations, and issue custom proctoring warnings. Experimental testing shows that the proctoring blocks successfully deter standard cheating methods, while the semantic grader exhibits high alignment with human-marked scores, proving both secure and operationally feasible.")

    # ==================== CHAPTER 1: INTRODUCTION ====================
    print("Writing Chapter 1 (Introduction)...")
    add_chapter_heading(doc, "1 Introduction")
    
    add_section_heading(doc, "1.1 Project Profile")
    add_paragraph(doc, "The development of a secure, proctored academic system requires a detailed parameters grid. The table below represents the complete profile parameters for the Online Examination System developed under Swaminarayan University:")
    
    tbl_profile = create_styled_table(doc, 1, 2, ["PARAMETER", "PROJECT DETAILS"], [2.5, 4.5])
    profile_data = [
        ("Project Title", "ONLINE-EXAMINATION-SYSTEM"),
        ("Developer Name", "Mayur Ramavat (Enrollment No: SUK230704CA062)"),
        ("Faculty Guide", "Mr. Dipesh Dave (Assistant Professor, Faculty of IT & CS)"),
        ("Academic Institution", "Swaminarayan University, Kalol, Gujarat"),
        ("Degree Course", "Bachelor of Computer Applications (B.C.A)"),
        ("Academic Year", "2025-2026"),
        ("Frontend Technologies", "HTML5, CSS3, JavaScript (ES6+), Web Audio API"),
        ("Backend Environment", "PHP 8.x, local Apache Server"),
        ("Database Engine", "MySQL (InnoDB Relational Storage)"),
        ("Semantic Grader", "Python 3.x, SequenceMatcher, String NLP overlap scripts"),
        ("Security Features", "Tab-switch detection, copy-paste disabled, bcrypt encryption, PDO prepared statements")
    ]
    for param, det in profile_data:
        add_table_row(tbl_profile, [param, det])
        
    add_section_heading(doc, "1.2 Project Summary")
    add_paragraph(doc, "The Online Examination System is an intelligent, secure web portal that automates student assessments, logs security violations, and grades student answers using a blend of deterministic checks and natural language overlap matching. The system consists of two primary portals: the student examination panel and the administrative proctor dashboard. Students login, view active examinations assigned by the university, and take exams with real-time countdown timers. The browser locks down input actions to prevent academic dishonesty. The system acts as a complete replacement for manual offline testing environments.")
    add_paragraph(doc, "The backend PHP architecture records draft responses in database tables in real-time, providing resistance against power cuts or network drops. Upon submission, the PHP application launches the Python grading script via secure command shell execution. The Python script pulls the student's submission, computes marks for multiple-choice questions by absolute matches, and evaluates descriptive answers by calculating a similarity coefficient compared to model answers. Supervisors monitor the entire process from a central proctor board that reports cheating logs instantly, ensuring credibility.")
    add_paragraph(doc, "The development of this software is aimed at improving both the administrative efficiency of universities and the integrity of online certifications. By providing automatic grading of descriptive answers, instructors are freed from tedious marking cycles, allowing them to focus on pedagogical improvements. At the same time, the browser lockdown controls guarantee that scores obtained are a true reflection of student knowledge, eliminating window-switching and copy-paste plagiarism.")

    add_section_heading(doc, "1.3 purpose: Goals & Objectives")
    add_paragraph(doc, "The primary goals and objectives of the developed system are laid out across distinct technological targets:")
    add_paragraph(doc, "1. Automated Descriptive Grading: Standard automated grading systems are limited to multiple-choice sheets because evaluating written sentences is highly complex. One of our main goals is to create a semantic grading module in Python that evaluates descriptive sentences by matching core keywords and calculating sentence similarity against instructor model keys. This permits subjective testing at scale without human graders. By automating descriptive grading, the grading workload for large student cohorts is reduced from days to a few seconds, while ensuring that the scores are objectively computed.")
    add_paragraph(doc, "2. Real-Time Proctoring and Anti-Cheating Controls: E-learning portals often suffer from a lack of integrity, as students can easily open secondary tabs to search for answers using web search engines. This project implements window blur and page visibility listeners that count how many times a candidate leaves the exam focus. If a candidate attempts to switch tabs or minimize the browser window to search for answers, the system immediately logs the incident and triggers a warning notification. If the candidate exceeds the threshold of five violations, the exam session is locked and force-submitted, protecting the assessment's integrity.")
    add_paragraph(doc, "3. Intercepting Input and Plagiarism Actions: Copying question text directly into search engines or pasting pre-composed essays is a major cheating threat. The system suppresses mouse right-clicks (context menu) and disables standard keyboard hooks (Ctrl+C, Ctrl+V, Ctrl+X, Ctrl+A). Text area fields also suppress drop events, forcing candidates to compose answers manually in the editor. This ensures that all submitted answers are written in real-time during the examination session, and not plagiarized from external sources.")
    add_paragraph(doc, "4. Real-time Supervisor Console: Invigilators need live visibility of candidates. The proctor dashboard displays cards of students writing exams. The cards update every 5 seconds using AJAX, displaying student names, enrollment numbers, active timers, tab switch statistics, and copy-paste attempts. Supervisors can monitor active sessions and send warning messages that display as overlays on the student's screen. If a student is flagged as suspicious, the supervisor can broadcast a custom warning alert, accompanied by an audio beep generated on the client browser.")
    add_paragraph(doc, "5. Fail-Safe Progress Recovery: If a student experiences a computer crash or power cut, their progress must be preserved. The student portal saves draft states to localStorage on keypress, and sends background AJAX requests to the server every 30 seconds. This double-layer backup guarantees that no more than 30 seconds of work is lost, allowing students to resume their exam immediately after rebooting. The system recovers their state automatically without requiring administrative intervention, ensuring a smooth exam experience.")

    add_section_heading(doc, "1.4 General Description")
    add_paragraph(doc, "The general application workflow is structured around four user types and their specific permissions. Each user interaction has been designed to maximize security and minimize latency:")
    add_paragraph(doc, "1. Visitor Module: This is the public portal. Visitors can view the university's main home page, read instructions about the proctored exam system, and access registration forms. They are blocked from accessing active exams, database panels, or the proctor monitoring console. If a visitor attempts to access restricted pages, PHP session filters intercept the request and redirect them to the login screen. This prevents unauthorized users from accessing active exam sessions.")
    add_paragraph(doc, "2. Student Module: Once logged in, students see their active dashboard. This panel lists active exams, upcoming dates, and completed results. Upon starting an exam, the student enters a proctored web layout. The top navigation bar displays their profile and a connection status indicator. The main body displays one question at a time to prevent cognitive overload. A sidebar palette displays numbers for all questions, indicating answered, unanswered, active, and flagged states. The browser records blur events, warning the student when they switch tabs. If warnings exceed five, the exam is locked and submitted automatically.")
    add_paragraph(doc, "3. Administrator Module: Admins manage academic parameters. They define exams, configure subjects, specify durations, and add questions. For each question, they set allocated points. MCQs require options A, B, C, D, and a correct key. Descriptive questions require a detailed model answer that represents the ideal response. Admins can view finalized student exams and review descriptive scores side-by-side with similarity indices, allowing them to verify automated grades.")
    add_paragraph(doc, "4. Supervisor Live Proctor Console: During active exams, supervisors use the proctor console to monitor the session. The page automatically fetches candidate details every 5 seconds, displaying them on color-coded cards. Clean sessions are green, suspicious sessions with warnings are orange, and terminated sessions are red. Supervisors can click a button to send a warning message (e.g. 'Please focus on your screen'). This message displays as an overlay on the student's screen, accompanied by an audio beep generated using the browser's Web Audio API.")

    # ==================== CHAPTER 2: SYSTEM ANALYSIS ====================
    print("Writing Chapter 2 (System Analysis)...")
    add_chapter_heading(doc, "2 System Analysis")
    
    add_section_heading(doc, "2.1 Study of Current System")
    add_paragraph(doc, "The current evaluation models in most universities can be divided into traditional paper assessments and basic online platforms. Traditional assessments require students to write answers on physical paper inside exam halls, under manual supervision. Instructors then collect and grade the papers. While traditional exams are secure, grading hundreds of scripts is extremely slow, and manual score calculations are prone to errors. This manual workflow creates significant administrative overhead, especially during end-of-semester assessments when faculty members must grade hundreds of multi-page answer sheets, leading to delays in publishing results.")
    add_paragraph(doc, "With the rise of e-learning, universities have adopted platforms like Moodle, Canvas, and Google Forms. These systems deliver MCQs online, grading them instantly by checking student choices against a key array. However, these basic portals lack security controls. Students can open secondary tabs, search for answers online, or copy-paste text between applications. Additionally, these platforms cannot grade descriptive answers, forcing instructors to manually read and grade written responses. This limitation restricts online exams to basic MCQ questionnaires, which fail to evaluate a student's analytical and subjective expression skills.")
    add_paragraph(doc, "The lack of security in basic online platforms leads to high rates of academic dishonesty. Because students are not monitored, they can search for answers on other websites or copy-paste pre-written text. This undermines the credibility of online certifications. Furthermore, manually grading hundreds of descriptive essays is highly tedious for instructors. A system that combines secure browser controls with automated descriptive grading is necessary to ensure credible online assessments.")

    add_section_heading(doc, "2.2 Problems of Current System")
    add_paragraph(doc, "An analysis of existing online and offline examination systems reveals several critical problems:")
    add_paragraph(doc, "1. Time-Consuming Descriptive Grading: Grading written answers requires significant time and effort from instructors. For large classes, manual grading causes long delays in publishing results. Additionally, grading subjective answers is prone to human fatigue and bias, where the same response may receive different marks depending on the grader's focus or mood. This subjective variation creates inconsistencies in student assessments and increases student complaints regarding grading fairness.")
    add_paragraph(doc, "2. Vulnerability to Online Cheating: Basic online exam portals do not restrict browser controls. Students can open additional tabs, use search engines, or share questions on messaging apps. Because the exam page remains open, students can easily search for answers on secondary screens without being logged or penalized. This lack of proctoring controls invalidates online exam scores, as candidates can easily find answers to factual and coding questions online.")
    add_paragraph(doc, "3. Copy-Paste Plagiarism: In standard text areas, students can copy-paste text from external files or AI tools. This allows candidates to copy-paste pre-written essays directly into the answer fields, bypassing the requirement to write the response during the examination session. This invalidates written assessments, as there is no way to verify if the student composed the text themselves during the exam.")
    add_paragraph(doc, "4. Lack of Progress Recovery: In many online test systems, if a student's computer reboots, their internet connection drops, or their browser closes unexpectedly, their progress is lost. This forces them to restart the exam, causing frustration and administrative delays. It also increases server load, as multiple students must restart their sessions simultaneously, creating additional connections.")
    add_paragraph(doc, "5. Lack of Real-Time Supervision: Online exams are often unsupervised. Invigilators cannot see student activities in real-time or log browser violations. This lack of visibility makes it difficult for administrators to verify the integrity of the examination session, as there are no audit logs to verify candidate compliance.")

    add_section_heading(doc, "2.3 Requirements of New System")
    add_paragraph(doc, "To resolve the limitations of the current system, the new Online Examination System implements several core technical requirements:")
    add_paragraph(doc, "1. Combined Grading Engine: The system must grade MCQs automatically using exact key matching, and grade descriptive answers semantically. The Python grader uses sequence matching and keyword overlap to evaluate subjective text, awarding marks based on how closely the response matches the key concepts in the model answer. This allows subjective assessments to be conducted and graded automatically, providing immediate feedback to students.")
    add_paragraph(doc, "2. Browser lockdown controls: The student portal must disable right-clicks (context menu) and block keyboard copy-paste shortcuts (Ctrl+C, Ctrl+V, Ctrl+X, Ctrl+A). This stops candidates from copy-pasting prepared materials into answer fields or copying questions for external searches. The browser must also block text dragging and dropping to ensure all text is composed within the editor.")
    add_paragraph(doc, "3. Tab Focus Monitoring: The application must log when a student leaves the exam focus. Window blur events increment the tab switch count. If a student switches tabs five times, the exam must lock automatically, submit the current answers, and redirect the student to the dashboard. This prevents students from searching for answers on secondary browser windows or applications.")
    add_paragraph(doc, "4. Live Proctoring Console: The admin panel must provide a live monitoring dashboard that refreshes every 5 seconds. The console displays active student statistics, including tab switch counts and copy-paste attempts. Admins can write warning messages that display as overlays on the student's screen, accompanied by an audio tone to alert the candidate.")
    add_paragraph(doc, "5. Double-Layer Progress Recovery: The student interface must save answer states to localStorage on every keypress, and send background draft saves to the server every 30 seconds. This ensures that progress is preserved in case of power or network failures, allowing students to resume their exam session immediately after rebooting.")

    add_section_heading(doc, "2.4 Feasibility Study")
    add_paragraph(doc, "Before starting development, a feasibility study was conducted to evaluate the project across five criteria:")
    add_paragraph(doc, "1. Technical Feasibility: The tech stack uses standard open-source technologies: PHP, MySQL, and JavaScript for the web application, and Python for the grading engine. These technologies are widely documented and run on typical server environments. Python's built-in libraries (SequenceMatcher) evaluate strings efficiently without requiring high-end GPU resources, making the system technically feasible. The integration of Python and PHP using system-level execution commands is stable and well-supported.")
    add_paragraph(doc, "2. Economic Feasibility: The system uses open-source tools, eliminating software license costs. The hardware requirements are minimal, as the system does not require expensive GPU/LLM models for grading. The operational cost savings from reduced paper printing, automated grading, and fewer manual proctors outweigh the development costs, making the project economically viable. The university can deploy the application on standard web hosting servers.")
    add_paragraph(doc, "3. Operational Feasibility: The user interfaces are designed using vanilla CSS for fast loading times. Students do not need to install additional software, as the proctoring functions run directly inside modern web browsers. Administrators can manage questions and monitor students using a simple web dashboard, ensuring high operational usability. Training requirements for administrators are minimal, as the interfaces are designed with intuitive navigation controls.")
    add_paragraph(doc, "4. Legal Feasibility: The system collects only student emails, enrollment numbers, and browser focus logs. No personal browser history or visual recordings are captured. This respects user privacy and complies with data protection regulations, ensuring the system is legally feasible. The system handles all user data in compliance with standard privacy policies.")
    add_paragraph(doc, "5. Social Feasibility: The system ensures a fair assessment environment by preventing cheating. This social benefit builds trust in online certifications. The automated grading engine also ensures objective scoring, which reduces student anxiety regarding evaluator bias, making the platform socially acceptable. By providing a secure and objective testing platform, the system supports modern online certification standards.")

    # ==================== CHAPTER 3: REQUIREMENT STUDY ====================
    print("Writing Chapter 3 (Requirement Study)...")
    add_chapter_heading(doc, "3 System Requirement Study")
    
    add_section_heading(doc, "3.1 User Characteristics")
    add_paragraph(doc, "The system interacts with three distinct user roles, each with specific access permissions and functional workflows:")
    
    add_paragraph(doc, "3.1.1 Visitor: Any guest user navigating to the landing page. Visitors can view the university's main home page, read instructions about the proctored exam system, and access student registration links. They have no access to exam session pages or administrative control panels. If a visitor attempts to access restricted pages, PHP session filters intercept the request and redirect them to the login screen. This prevents unauthorized users from accessing active exam sessions.")
    add_paragraph(doc, "3.1.2 User: Registered students who log in to their dashboard. Students have basic technical skills to use web browsers. They view active exams, start assessments, answer questions, navigate using a sidebar palette, receive proctoring warning dialogues, and view finalized scores after grading. They are expected to be familiar with basic browser controls and online test formats.")
    add_paragraph(doc, "3.1.3 Admin: University faculty members or administrative staff who manage exams. Admins create exam definitions, set duration timers, add questions, specify correct MCQ keys and descriptive model answers, monitor student exams in real-time, view cheating logs, issue warnings, and reset student exam states. They require standard administrative access and are trained in managing exam parameters.")

    add_section_heading(doc, "3.2 Software & Hardware Requirement")
    add_subsection_heading(doc, "3.2.1 Minimum Hardware Requirement")
    add_paragraph(doc, "The following hardware configurations are required to deploy and run the system:")
    
    tbl_hw = create_styled_table(doc, 1, 3, ["COMPONENT", "CLIENT MACHINE (STUDENT)", "SERVER MACHINE"], [2.0, 2.5, 2.5])
    add_table_row(tbl_hw, ["CPU Processor", "Intel Core i3 or equivalent (2.0 GHz+)", "Intel Xeon or Modern multi-core CPU (2.4 GHz+)"])
    add_table_row(tbl_hw, ["System Memory (RAM)", "4 GB DDR3 / DDR4 minimum", "8 GB DDR4 minimum (16 GB recommended)"])
    add_table_row(tbl_hw, ["Hard Disk Space", "500 MB free space (for browser cache)", "20 GB free SSD space (for DB and file storage)"])
    add_table_row(tbl_hw, ["Network Interface", "Standard Wi-Fi / Ethernet (1 Mbps+)", "High-speed Ethernet (100 Mbps+ public link)"])
    add_table_row(tbl_hw, ["Peripherals", "Keyboard, Mouse (Webcam optional for future)", "Standard server console, backup storage arrays"])
    
    add_subsection_heading(doc, "3.2.2 Minimum Software Requirement")
    add_paragraph(doc, "The following software systems are required to configure and run the application development environment:")
    
    tbl_sw = create_styled_table(doc, 1, 3, ["SOFTWARE TYPE", "DEVELOPMENT TOOL", "MINIMUM VERSION"], [2.0, 3.5, 1.5])
    add_table_row(tbl_sw, ["Operating System", "Microsoft Windows 10/11 or Ubuntu Linux", "OS Independent"])
    add_table_row(tbl_sw, ["Web Server", "Apache Web Server (included in XAMPP)", "Apache 2.4 / Nginx"])
    add_table_row(tbl_sw, ["Database Engine", "MySQL Database", "MySQL 5.7 / MariaDB 10.4"])
    add_table_row(tbl_sw, ["Backend Language", "PHP (Hypertext Preprocessor)", "PHP 8.0 or higher"])
    add_table_row(tbl_sw, ["AI/Semantic Engine", "Python Interpreter", "Python 3.8 or higher"])
    add_table_row(tbl_sw, ["Required Python Packages", "mysql-connector-python", "latest stable"])
    add_table_row(tbl_sw, ["Web Browser", "Google Chrome, Mozilla Firefox, Microsoft Edge", "Modern standard browser"])
    add_table_row(tbl_sw, ["Text Editor", "Visual Studio Code or Notepad++", "Latest Version"])

    add_section_heading(doc, "3.3 Functional and Non functional requirement")
    add_subsection_heading(doc, "3.3.1 Functional Requirements")
    add_paragraph(doc, "The functional requirements define the core operations that the system must perform:")
    add_bullet_point(doc, "The system must verify student credentials and establish session tokens, preventing unauthorized logins.", "Authentication:")
    add_bullet_point(doc, "The system must load exam metadata, duration, and corresponding question lists without displaying answer keys in the HTML DOM.", "Exam Delivery:")
    add_bullet_point(doc, "The client interface must intercept right-clicks, copy, paste, select, and cut events, preventing these actions within the exam container.", "Input Restriction:")
    add_bullet_point(doc, "The client must detect when the browser window loses focus, increment the tab switch count, and alert the student.", "Tab Monitoring:")
    add_bullet_point(doc, "The system must lock the exam and auto-submit current draft answers once the tab switch count reaches five.", "Forced Lockout:")
    add_bullet_point(doc, "Every 30 seconds, student answers must be sent to the server as a background draft, and every 10 seconds, a heartbeat must log status.", "Autosave & Status Heartbeat:")
    add_bullet_point(doc, "On submission, the system must invoke the Python grading script, compute scores, update database fields, and change exam status.", "AI Auto Grading:")
    add_bullet_point(doc, "The admin panel must provide a live dashboard with active student proctor metrics and manual warning fields.", "Proctor Console:")
    
    add_subsection_heading(doc, "3.3.2 Non-functional Requirements")
    add_paragraph(doc, "The non-functional requirements specify constraints on how the system performs:")
    add_bullet_point(doc, "The system must load exam dashboards and submit answers in less than 2 seconds over standard broadband networks.", "Performance:")
    add_bullet_point(doc, "The platform must secure user passwords using bcrypt hashing and protect all SQL queries against injections using prepared statements.", "Security:")
    add_bullet_point(doc, "The student exam interface must adapt to mobile devices, tablets, and desktop displays using responsive layouts.", "Usability & Responsiveness:")
    add_bullet_point(doc, "The system must handle multiple concurrent student submissions, utilizing MySQL transaction tables and efficient Python string libraries.", "Scalability:")
    add_bullet_point(doc, "The PHP and Python code should be structured in distinct directories (config, engine, api, assets) for easy maintainability.", "Modularity:")

    # ==================== CHAPTER 4: SYSTEM DESIGN ====================
    print("Writing Chapter 4 (System Design)...")
    add_chapter_heading(doc, "4 System Design")
    
    add_section_heading(doc, "4.1 System Diagram")
    add_paragraph(doc, "The system is structured around a Three-Tier Client-Server Architecture. The client tier comprises the student's browser rendering HTML, CSS, and vanilla JS. The logic tier consists of the Apache web server running PHP scripts, which handle authentication, exam flows, and proctor logging. The database tier uses a MySQL server to maintain relational schemas, while the Python script executes on the server side, acting as an offline NLP evaluator. The architecture layout is detailed below:")
    add_diagram(doc, "system_architecture.png", "System Architecture (Three-Tier Diagram)")
    add_paragraph(doc, "The separation of concerns between PHP and Python ensures that the web application remains responsive. While PHP handles user routing and database connection management, the heavy calculation of sequence similarity is offloaded to the Python interpreter. The command line execution model keeps the interface fast, since the grader executes only when the student submits the exam. The PHP script initiates the grader using shell commands, passing the unique exam session ID as a parameter.")
    
    add_section_heading(doc, "4.2 Activity Diagram")
    add_paragraph(doc, "The Activity Diagram maps the step-by-step workflow of a student exam session. After logging in, the student selects an active exam. The exam page initializes timers and listeners. While attempting questions, if the student switches tabs, the blur event increases the tab switch counter. If the count reaches 5, the exam is locked and submitted. Otherwise, the student can submit manually. The submission triggers the Python grader, which updates scores in the database, concluding the process:")
    add_diagram(doc, "activity_diagram.png", "Student Assessment & Proctoring Activity Flow")
    add_paragraph(doc, "The client-side scripts handle user input in real-time, preventing cheating before it happens. At the same time, the server monitors connection state through heartbeat queries. If a student attempts to block the JavaScript file, the server records a heartbeat loss. The proctor dashboard flags this, allowing the supervisor to issue warnings or terminate the exam session manually. The active interface updates immediately when warning payloads are received in the heartbeat response.")
    
    add_section_heading(doc, "4.3 Data Flow Diagram")
    add_paragraph(doc, "The Data Flow Diagram (Level 1) details how data travels between external entities (Student and Admin), internal processes, and the database tables. Process 1.0 validates user logins. Process 2.0 manages active exams, writing progress and proctor violation counts to student_exams and student_answers. Process 3.0 allows admins to define exams and manage questions in database tables. Process 4.0 (the Python Grader) reads questions and student answers, compares them to model answers, and writes scores and feedback to the database:")
    add_diagram(doc, "dfd_diagram.png", "Data Flow Diagram (DFD Level 1)")
    add_paragraph(doc, "Level 0 context diagrams represent the system as a single process block with inputs and outputs from Student and Admin. Level 1 diagrams break this block into four distinct processes, showing data flows to database tables like `student_exams` and `questions`. Level 2 diagrams detail specific subprocesses, such as the grading loop and proctor verification flows. Each data store in the diagram represents a physical table in our normalized relational database.")
    
    add_section_heading(doc, "4.4 Use Case Diagram")
    add_paragraph(doc, "The Use Case Diagram displays actor-system interactions. The Student (user) interacts with registration, logins, exam page displays, and auto-saves. The Admin (proctor) manages exams and questions, monitors live student activities, sends proctor warnings, and triggers the Python AI grading script:")
    add_diagram(doc, "use_case_diagram.png", "Use Case Interaction Model")
    add_paragraph(doc, "Each use case represents a distinct functional module. For instance, the 'Attempt Online Exam' use case requires the 'Verify Session Token' use case. The proctoring use cases, like 'Send Proctor Warning' and 'Monitor Live Exam', interact with the active exam session database, creating a secure assessment environment. These use cases define our test plans and functional design boundaries.")
    
    add_section_heading(doc, "4.5 Data Modelling")
    add_subsection_heading(doc, "4.5.1 E-R Diagram")
    add_paragraph(doc, "The Entity-Relationship (E-R) Diagram represents the database structure, showing entities and cardinialities. Admins create Exams (1:N relation). Exams contain Questions (1:N relation). Students take Exams, creating a Student Exams relation. Student Exams record Student Answers, which are linked to Questions, creating a robust relational model:")
    add_diagram(doc, "er_diagram.png", "Entity-Relationship (E-R) Database Schema")
    add_paragraph(doc, "The ER model is normalized to Third Normal Form (3NF) to prevent redundancy. For instance, student details are stored once in `students`, while exam sessions are logged in `student_exams`. The `student_answers` table records answers to specific questions, linking them to student exams via foreign keys. Primary and foreign key indexes are configured across all join vectors to ensure high-speed lookups.")
    
    add_subsection_heading(doc, "4.5.2 Data Dictionary")
    add_paragraph(doc, "The following tables define the relational database schema, columns, data types, keys, and descriptions for the Swaminarayan University Online Exam database (`su_exam_db`):")
    
    # 1. Admins Table
    add_paragraph(doc, "Table 4.1: Admins Table Schema")
    t_adm = create_styled_table(doc, 1, 5, ["COLUMN NAME", "DATA TYPE", "CONSTRAINTS", "DEFAULT", "DESCRIPTION"], [1.5, 1.2, 1.2, 1.0, 2.1])
    add_table_row(t_adm, ["id", "INT", "PK, Auto-Increment", "None", "Unique identifier for administrator account"])
    add_table_row(t_adm, ["full_name", "VARCHAR(100)", "NOT NULL", "None", "First and last name of admin"])
    add_table_row(t_adm, ["email", "VARCHAR(100)", "UNIQUE, NOT NULL", "None", "Email used for admin login credentials"])
    add_table_row(t_adm, ["password", "VARCHAR(255)", "NOT NULL", "None", "Bcrypt hashed password string"])
    add_table_row(t_adm, ["contact_no", "VARCHAR(20)", "NULL", "NULL", "Mobile/telephone number of admin"])
    add_table_row(t_adm, ["created_at", "TIMESTAMP", "NOT NULL", "CURRENT_TIMESTAMP", "Date and time of record insertion"])
    
    add_paragraph(doc, "Detailed Column Analysis (Admins Table):")
    add_bullet_point(doc, "The primary key identifier, configured as an auto-increment integer, ensuring that every administrator profile is cataloged uniquely.", "1. id (INT):")
    add_bullet_point(doc, "Stores the supervisor's official name, capped at 100 characters. Used for welcome banners and proctor alerts.", "2. full_name (VARCHAR):")
    add_bullet_point(doc, "The email address serves as the login username. It features a unique constraint to prevent registering multiple admin profiles with the same mail record.", "3. email (VARCHAR):")
    add_bullet_point(doc, "A secure password field of length 255 to store the Bcrypt output of the password, preventing text exposure.", "4. password (VARCHAR):")
    
    # 2. Students Table
    add_paragraph(doc, "Table 4.2: Students Table Schema")
    t_std = create_styled_table(doc, 1, 5, ["COLUMN NAME", "DATA TYPE", "CONSTRAINTS", "DEFAULT", "DESCRIPTION"], [1.5, 1.2, 1.2, 1.0, 2.1])
    add_table_row(t_std, ["id", "INT", "PK, Auto-Increment", "None", "Unique identifier for student record"])
    add_table_row(t_std, ["enrollment_no", "VARCHAR(50)", "UNIQUE, NOT NULL", "None", "University enrollment registration number"])
    add_table_row(t_std, ["full_name", "VARCHAR(100)", "NOT NULL", "None", "Full name of the student"])
    add_table_row(t_std, ["email", "VARCHAR(100)", "UNIQUE, NOT NULL", "None", "University student email address"])
    add_table_row(t_std, ["password", "VARCHAR(255)", "NOT NULL", "None", "Bcrypt hashed password string"])
    add_table_row(t_std, ["contact_no", "VARCHAR(20)", "NULL", "NULL", "Mobile number for SMS communications"])
    add_table_row(t_std, ["profile_photo", "VARCHAR(255)", "NULL", "NULL", "File path to student avatar image"])
    add_table_row(t_std, ["created_at", "TIMESTAMP", "NOT NULL", "CURRENT_TIMESTAMP", "Timestamp of registration"])
    
    add_paragraph(doc, "Detailed Column Analysis (Students Table):")
    add_bullet_point(doc, "Unique integer matching each registered candidate, used for foreign key joins in session tables.", "1. id (INT):")
    add_bullet_point(doc, "The university registration code (e.g. SUK230704CA062). Features a unique key constraint to verify students during authentication.", "2. enrollment_no (VARCHAR):")
    add_bullet_point(doc, "Stores the student's full name, used for display banners and grading worksheets.", "3. full_name (VARCHAR):")
    add_bullet_point(doc, "Hashed character array containing the candidate's secret password, protecting student records.", "4. password (VARCHAR):")
    
    # 3. Exams Table
    add_paragraph(doc, "Table 4.3: Exams Table Schema")
    t_exm = create_styled_table(doc, 1, 5, ["COLUMN NAME", "DATA TYPE", "CONSTRAINTS", "DEFAULT", "DESCRIPTION"], [1.5, 1.2, 1.2, 1.0, 2.1])
    add_table_row(t_exm, ["id", "INT", "PK, Auto-Increment", "None", "Unique identifier for examination"])
    add_table_row(t_exm, ["title", "VARCHAR(150)", "NOT NULL", "None", "Title/Name of the examination"])
    add_table_row(t_exm, ["subject", "VARCHAR(100)", "NOT NULL", "None", "Academic subject/course boundary"])
    add_table_row(t_exm, ["duration_minutes", "INT", "NOT NULL", "None", "Total allowed time in minutes"])
    add_table_row(t_exm, ["created_by", "INT", "FK -> admins(id)", "NULL", "Admin id who defined this exam"])
    add_table_row(t_exm, ["status", "ENUM('draft','active','completed')", "NOT NULL", "'draft'", "Exam current status"])
    add_table_row(t_exm, ["created_at", "TIMESTAMP", "NOT NULL", "CURRENT_TIMESTAMP", "Creation timestamp"])
    
    add_paragraph(doc, "Detailed Column Analysis (Exams Table):")
    add_bullet_point(doc, "The unique identifier for the examination, referenced by the questions and student exam tables.", "1. id (INT):")
    add_bullet_point(doc, "The academic subject of the exam (e.g. Computer Science).", "2. subject (VARCHAR):")
    add_bullet_point(doc, "The allowed duration in minutes (e.g., 30, 60, or 120 minutes), used to initialize the client countdown timer.", "3. duration_minutes (INT):")
    add_bullet_point(doc, "An enum displaying the current state: 'draft' (hidden), 'active' (visible and testable), or 'completed' (locked).", "4. status (ENUM):")
    
    # 4. Questions Table
    add_paragraph(doc, "Table 4.4: Questions Table Schema")
    t_qst = create_styled_table(doc, 1, 5, ["COLUMN NAME", "DATA TYPE", "CONSTRAINTS", "DEFAULT", "DESCRIPTION"], [1.5, 1.2, 1.2, 1.0, 2.1])
    add_table_row(t_qst, ["id", "INT", "PK, Auto-Increment", "None", "Unique identifier for question"])
    add_table_row(t_qst, ["exam_id", "INT", "FK -> exams(id)", "None", "Associated exam record ID"])
    add_table_row(t_qst, ["question_text", "TEXT", "NOT NULL", "None", "Body text of the question"])
    add_table_row(t_qst, ["type", "ENUM('mcq','descriptive')", "NOT NULL", "None", "Question category (MCQ or subjective text)"])
    add_table_row(t_qst, ["option_a", "VARCHAR(255)", "NULL", "NULL", "Text option A (MCQ only)"])
    add_table_row(t_qst, ["option_b", "VARCHAR(255)", "NULL", "NULL", "Text option B (MCQ only)"])
    add_table_row(t_qst, ["option_c", "VARCHAR(255)", "NULL", "NULL", "Text option C (MCQ only)"])
    add_table_row(t_qst, ["option_d", "VARCHAR(255)", "NULL", "NULL", "Text option D (MCQ only)"])
    add_table_row(t_qst, ["correct_option", "CHAR(1)", "NULL", "NULL", "Correct option letter A, B, C, or D (MCQ only)"])
    add_table_row(t_qst, ["model_answer", "TEXT", "NULL", "NULL", "Expected answer text for similarity grader"])
    add_table_row(t_qst, ["points", "INT", "NOT NULL", "1", "Points allocated to this question"])
    
    add_paragraph(doc, "Detailed Column Analysis (Questions Table):")
    add_bullet_point(doc, "Foreign key connecting the question to its parent exam profile. Cascade delete is configured so deleting an exam removes all its questions.", "1. exam_id (INT):")
    add_bullet_point(doc, "Categorizes the question format: 'mcq' (options rendered) or 'descriptive' (large textarea rendered).", "2. type (ENUM):")
    add_bullet_point(doc, "Holds the correct option (A, B, C, or D) for automated MCQ grading.", "3. correct_option (CHAR):")
    add_bullet_point(doc, "Stores the expected answer string used by the Python grading script to calculate descriptive score similarity.", "4. model_answer (TEXT):")
    
    # 5. Student Exams Table
    add_paragraph(doc, "Table 4.5: Student Exams Table Schema")
    t_sex = create_styled_table(doc, 1, 5, ["COLUMN NAME", "DATA TYPE", "CONSTRAINTS", "DEFAULT", "DESCRIPTION"], [1.5, 1.2, 1.2, 1.0, 2.1])
    add_table_row(t_sex, ["id", "INT", "PK, Auto-Increment", "None", "Unique identifier for student exam session"])
    add_table_row(t_sex, ["student_id", "INT", "FK -> students(id)", "None", "ID of candidate writing exam"])
    add_table_row(t_sex, ["exam_id", "INT", "FK -> exams(id)", "None", "ID of exam being taken"])
    add_table_row(t_sex, ["status", "ENUM('started','submitted','graded')", "NOT NULL", "'started'", "Student exam session status"])
    add_table_row(t_sex, ["started_at", "TIMESTAMP", "NOT NULL", "CURRENT_TIMESTAMP", "Timestamp of starting the exam"])
    add_table_row(t_sex, ["submitted_at", "TIMESTAMP", "NULL", "NULL", "Timestamp of final submission"])
    add_table_row(t_sex, ["tab_switch_count", "INT", "NOT NULL", "0", "Total tab switch focus blur violations"])
    add_table_row(t_sex, ["copy_paste_count", "INT", "NOT NULL", "0", "Total intercepted keyboard copy/paste events"])
    add_table_row(t_sex, ["score", "DECIMAL(5,2)", "NOT NULL", "0.00", "Final graded score calculated by engine"])
    add_table_row(t_sex, ["total_possible_score", "INT", "NOT NULL", "0", "Max possible points for the exam"])
    add_table_row(t_sex, ["proctor_warning", "TEXT", "NULL", "NULL", "Warning text broadcasted by live proctor"])
    add_table_row(t_sex, ["last_active", "TIMESTAMP", "NULL", "NULL", "Timestamp of last heartbeat ping"])

    add_paragraph(doc, "Detailed Column Analysis (Student Exams Table):")
    add_bullet_point(doc, "An integer logging tab-switch events. If this count reaches 5, the exam is locked and submitted.", "1. tab_switch_count (INT):")
    add_bullet_point(doc, "An integer logging keyboard copy-paste attempts. Used on the proctor dashboard to identify suspicious activities.", "2. copy_paste_count (INT):")
    add_bullet_point(doc, "Holds the cumulative score of the student, calculated dynamically by the Python script.", "3. score (DECIMAL):")
    add_bullet_point(doc, "Contains warning text written by the proctor, fetched by the student's browser every 10 seconds.", "4. proctor_warning (TEXT):")
    add_bullet_point(doc, "Updated by the 10-second heartbeat to check if the student's browser is active.", "5. last_active (TIMESTAMP):")

    # 6. Student Answers Table
    add_paragraph(doc, "Table 4.6: Student Answers Table Schema")
    t_san = create_styled_table(doc, 1, 5, ["COLUMN NAME", "DATA TYPE", "CONSTRAINTS", "DEFAULT", "DESCRIPTION"], [1.5, 1.2, 1.2, 1.0, 2.1])
    add_table_row(t_san, ["id", "INT", "PK, Auto-Increment", "None", "Unique identifier for candidate answer row"])
    add_table_row(t_san, ["student_exam_id", "INT", "FK -> student_exams(id)", "None", "Linked exam session ID"])
    add_table_row(t_san, ["question_id", "INT", "FK -> questions(id)", "None", "Linked question ID"])
    add_table_row(t_san, ["student_answer", "TEXT", "NOT NULL", "None", "Candidate answer text or selected option"])
    add_table_row(t_san, ["marks_obtained", "DECIMAL(5,2)", "NOT NULL", "0.00", "Marks allocated by AI grader for this question"])
    add_table_row(t_san, ["auto_feedback", "TEXT", "NULL", "NULL", "Evaluation feedback and similarity index log"])

    add_paragraph(doc, "Detailed Column Analysis (Student Answers Table):")
    add_bullet_point(doc, "Foreign key connecting the answer row to the specific student exam session.", "1. student_exam_id (INT):")
    add_bullet_point(doc, "Foreign key connecting the response to the corresponding question in the questions table.", "2. question_id (INT):")
    add_bullet_point(doc, "Stores the student's response string (the selected MCQ letter or descriptive text).", "3. student_answer (TEXT):")
    add_bullet_point(doc, "The calculated points awarded to the candidate for this response.", "4. marks_obtained (DECIMAL):")
    add_bullet_point(doc, "Logs details of the grading process, such as keyword match counts and similarity indices.", "5. auto_feedback (TEXT):")

    # ==================== CHAPTER 5: PLANNING & SCHEDULING ====================
    print("Writing Chapter 5 (Planning & Scheduling)...")
    add_chapter_heading(doc, "5 Project Planning & Scheduling")
    
    add_section_heading(doc, "5.1 Process model")
    add_paragraph(doc, "The software development lifecycle (SDLC) followed for this project was the Agile Process Model. This methodology emphasizes iterative development, continuous feedback, and rapid adjustments. The project was broken down into three 2-week sprints, each resulting in an operational increment of the exam platform. Sprint 1 focused on basic student-admin authentication and question management database structures. Sprint 2 implemented the live examination browser layout, including the prevention of right-click and keyboard copying. Sprint 3 integrated the Python semantic grading engine and built the live proctor monitoring dashboard.")
    add_paragraph(doc, "Agile development allows for rapid adjustments based on feedback. For instance, when initial testing showed that students could bypass copy blocks by dragging text into fields, we quickly updated the code to suppress drop events on textareas. The daily meetings and sprint retrospectives kept development aligned with academic and security standards, ensuring a robust final application.")
    add_paragraph(doc, "By utilizing Scrum boards, we kept track of tasks and sprints. The product backlog list was continuously updated with features like security logs and warning tones. This iterative model allowed us to test each module before integrating it, reducing development risks and ensuring that the final application was highly stable and secure under concurrent user loads.")
    
    add_section_heading(doc, "5.2 Project Plan")
    add_paragraph(doc, "The project plan was structured across 12 weeks of development, detailing specific milestones and milestone phases:")
    add_paragraph(doc, "Weeks 1-2: Gathering requirements from the department. Designing the database tables, schemas, and relational keys. Modeling data flows and drawing UML diagrams. Researching visibility and focus APIs to verify proctoring controls.")
    add_paragraph(doc, "Weeks 3-4: Configuring the MySQL database. Coding PHP login and registration APIs. Implementing secure session handling and authentication tokens. Setting up the directory structures for assets, config, and backend engines.")
    add_paragraph(doc, "Weeks 5-6: Building the student dashboard and examination pages. Implementing timers, question palettes, and auto-save APIs. Configuring CSS styles with saffron and navy themes.")
    add_paragraph(doc, "Weeks 7-8: Scripting browser controls. Writing focus and visibility listeners to track tab changes. Suppressing right-clicks and copy-paste shortcuts. Designing the heartbeat polling API to sync student status every 10 seconds.")
    add_paragraph(doc, "Weeks 9-10: Coding the Python grading script. Connecting it to MySQL and testing similarity matching. Building admin dashboard pages, including active proctor monitor cards and warning forms.")
    add_paragraph(doc, "Weeks 11-12: System testing, bug fixes, and documentation. Running simulated exams to verify stability, security blocks, and grading accuracy. Completing the project report for submission.")

    add_section_heading(doc, "5.3 Schedule Representation")
    add_paragraph(doc, "The development task allocation and schedules are represented in the table below:")
    
    tbl_sched = create_styled_table(doc, 1, 4, ["TASK NAME", "START DATE", "END DATE", "DURATION (DAYS)"], [2.5, 1.5, 1.5, 1.5])
    add_table_row(tbl_sched, ["Requirements & UML Modeling", "2026-03-01", "2026-03-12", "12"])
    add_table_row(tbl_sched, ["DB Schema & Auth API Coding", "2026-03-13", "2026-03-24", "12"])
    add_table_row(tbl_sched, ["Exam Page UI & Timer Build", "2026-03-25", "2026-04-05", "12"])
    add_table_row(tbl_sched, ["Tab Switch & Keyboard Hook Security", "2026-04-06", "2026-04-17", "12"])
    add_table_row(tbl_sched, ["Python NLP Grader Development", "2026-04-18", "2026-04-30", "13"])
    add_table_row(tbl_sched, ["Admin Live Proctor Dashboard", "2026-05-01", "2026-05-12", "12"])
    add_table_row(tbl_sched, ["System Testing & Bug Resolution", "2026-05-13", "2026-05-24", "12"])

    # ==================== CHAPTER 6: RISK MANAGEMENT ====================
    print("Writing Chapter 6 (Risk Management)...")
    add_chapter_heading(doc, "6 Risk Management")
    
    add_section_heading(doc, "6.1 Risk Identification")
    add_paragraph(doc, "Identifying technical and operational risks is essential to ensure the reliability of the system under live exam loads. The following risks were identified during analysis:")
    add_paragraph(doc, "1. Client-side browser crash: A student's computer may crash or their browser may close unexpectedly mid-exam, threatening to lose their progress. This risk is common in regions with unstable power grids.")
    add_paragraph(doc, "2. JavaScript deactivation: Students may try to disable JavaScript in their browser settings to bypass the proctoring listeners, allowing them to switch tabs without warnings.")
    add_paragraph(doc, "3. Grading database errors: Connection drops between PHP and Python can cause grading tasks to fail, leaving exams ungraded or saving invalid results.")
    add_paragraph(doc, "4. High server load: Multiple concurrent submissions at the end of an exam can overload database connections, causing server timeouts and losing student scores.")
    add_paragraph(doc, "5. Incorrect similarity scores: The grading script may award low scores to correct answers that use different synonyms or sentence structures from the model answer.")

    add_section_heading(doc, "6.2 Risk Analysis")
    add_paragraph(doc, "The identified risks were analyzed using a Probability-Impact assessment matrix:")
    
    tbl_risk = create_styled_table(doc, 1, 4, ["RISK DESCRIPTION", "PROBABILITY", "IMPACT", "RISK LEVEL"], [2.5, 1.5, 1.5, 1.5])
    add_table_row(tbl_risk, ["Client machine reboot / crash", "Medium", "High", "Medium-High"])
    add_table_row(tbl_risk, ["JavaScript bypass attempt", "Low", "Critical", "Medium"])
    add_table_row(tbl_risk, ["Python execution error", "Low", "High", "Low-Medium"])
    add_table_row(tbl_risk, ["Database transaction lockout", "Medium", "High", "Medium-High"])
    add_table_row(tbl_risk, ["Incorrect similarity scoring", "Medium", "Medium", "Medium"])
    
    add_section_heading(doc, "6.3 Risk Planning")
    add_paragraph(doc, "To mitigate these risks, several controls were built into the software architecture:")
    add_paragraph(doc, "1. Double-Layer Progress Backups: The client interface saves progress to localStorage on keypress, and sends background draft saves to the server every 30 seconds. This ensures that no more than 30 seconds of work is lost, allowing students to resume their exam session immediately after rebooting.")
    add_paragraph(doc, "2. Server-side validation: The server validates all authentication tokens and sessions before accepting submissions. If a student disables JavaScript, the server stops receiving the required heartbeat pings, flagging their session as inactive on the proctor dashboard. This alerts the supervisor immediately.")
    add_paragraph(doc, "3. Synonyms logic: The Python grader calculates a hybrid similarity score, combining sequence matching with keyword overlap. By scaling scores so that a 75% match receives full marks, the grading engine accounts for minor differences in wording and synonyms. This reduces false scoring errors.")
    add_paragraph(doc, "4. DB connection pooling: XAMPP server configurations are optimized to support up to 500 concurrent connections, while tables use InnoDB transactions to prevent locks during simultaneous database writes. Database schemas are indexed to optimize queries.")

    # ==================== CHAPTER 7: SECURITY FEATURES ====================
    print("Writing Chapter 7 (Security Features)...")
    add_chapter_heading(doc, "7 Security Features")
    
    add_paragraph(doc, "Security is the defining parameter of the Swaminarayan University Online Exam portal. Several modern web protection strategies have been implemented to ensure exam credibility:")
    
    add_subsection_heading(doc, "7.1 Browser Window Focus Enforcement")
    add_paragraph(doc, "Using the browser's Document Object Model (DOM) APIs, the system binds event listeners to window focus, blur, and visibility changes. When a student switches tabs, opens an application like Discord, or switches screens, a blur event is logged. The system shows a red toast alert and records the count. Once the violation count exceeds five, the interface locks the candidate out and auto-submits current answers to secure the session's integrity.")
    add_paragraph(doc, "The blur detection logic uses JavaScript's `window.addEventListener('blur', ...)` event, which triggers when the page loses focus. The page visibility state is also monitored using `document.addEventListener('visibilitychange', ...)`, detecting when the tab is minimized or hidden. These events send AJAX payloads to `api/submit_exam.php`, updating the violation count in the database in real-time. By writing violation logs directly to the database, the system ensures that students cannot clear logs by reloading the page.")
    
    add_subsection_heading(doc, "7.2 Keyboard and Mouse Event Blockers")
    add_paragraph(doc, "To prevent students from copying question texts or pasting pre-composed paragraphs, the frontend script binds preventDefault() actions to standard mouse and keyboard hooks:")
    add_bullet_point(doc, "Right-click is disabled globally, preventing the copy/paste context menu from opening and stopping candidates from inspecting the page source.", "Context Menu Block:")
    add_bullet_point(doc, "Intercepts keyboard combinations (Ctrl+C, Ctrl+V, Ctrl+X, Ctrl+A, Ctrl+P) on the body level, blocking these actions and logging warning alerts.", "Shortcut Key Blocking:")
    add_bullet_point(doc, "Text area elements explicitly block paste and cut events, stopping copy-paste behaviors. Drag-and-drop events are also suppressed on textareas.", "Field-Level Input Restrictions:")
    
    add_subsection_heading(doc, "7.3 Secure Database & Data Transport")
    add_paragraph(doc, "All communications between the browser and the web server are processed securely. The following database security controls are active:")
    add_bullet_point(doc, "All database connections are established via PHP Data Objects (PDO) with prepared SQL parameter binds. This protects the system against SQL injection attempts by separating SQL command strings from variable parameters.", "PDO Prepared SQL Statements:")
    add_bullet_point(doc, "Student passwords are encrypted using bcrypt hashing in PHP, protecting user records against unauthorized database access. Hashing is processed using a cost factor of 10.", "Password Encryption via bcrypt:")
    add_bullet_point(doc, "Session values are validated on every script reload. Exams page links require valid, matching student_exam_id and user_id sessions to prevent session hijacking. The system automatically terminates sessions after 30 minutes of inactivity.", "Secure Sessions:")
    add_bullet_point(doc, "All text input parameters are sanitized using htmlspecialchars() before rendering, preventing cross-site scripting (XSS) attacks in descriptive fields.", "Input Sanitization:")

    # ==================== CHAPTER 8: TESTING ====================
    print("Writing Chapter 8 (Testing & Test Cases)...")
    add_chapter_heading(doc, "8 Testing")
    
    add_section_heading(doc, "8.1 Testing Strategy")
    add_paragraph(doc, "The testing strategy combined unit, integration, and system testing. Unit testing evaluated individual components, such as verifying bcrypt hashes and testing Python string clean functions. Integration testing checked API connectivity between frontend fetch() payloads and backend PHP scripts. System testing validated end-to-end functionality: a student logging in, taking an exam, triggering tab switches, and verifying the AI grading results in the database.")
    add_paragraph(doc, "Stress testing was also conducted, simulating multiple users submitting tests simultaneously. This confirmed that the PHP script and Python grader execute efficiently without locking the database or causing server crashes. The grading algorithms were validated using real subjective responses, showing a high correlation with manual marks. Performance evaluations checked query response times under high-load conditions.")
    
    add_section_heading(doc, "8.2 Testing Methods")
    add_paragraph(doc, "The system underwent extensive black-box testing. The table below lists 25 detailed test cases:")
    
    # 25 Detailed Test Cases Table
    tbl_test = create_styled_table(doc, 1, 5, ["ID", "TEST SCENARIO", "INPUT DATA", "EXPECTED OUTCOME", "STATUS"], [0.6, 2.2, 1.2, 2.2, 0.8])
    test_cases_data = [
        ("TC01", "Student login verification", "Valid email & password", "Success, redirects to dashboard", "PASS"),
        ("TC02", "Invalid student login", "Wrong password input", "Error shown, redirects blocked", "PASS"),
        ("TC03", "Admin login validation", "Valid admin credentials", "Redirects to admin dashboard", "PASS"),
        ("TC04", "Invalid admin login", "Wrong admin credentials", "Authentication fails with alert", "PASS"),
        ("TC05", "Right-click context block", "Mouse right-click action", "Menu blocked, warning toast shown", "PASS"),
        ("TC06", "Ctrl+C shortcut lock", "Press Ctrl+C on question text", "Copy action blocked, alert logged", "PASS"),
        ("TC07", "Ctrl+V shortcut lock", "Press Ctrl+V in descriptive box", "Paste action blocked, alert logged", "PASS"),
        ("TC08", "Drag and drop blocking", "Drag external text into textarea", "Drop event canceled, no text added", "PASS"),
        ("TC09", "Tab Switch focus-blur count", "Alt+Tab to change window", "Warning toast displays switch count", "PASS"),
        ("TC10", "Forced submission lockout", "Alt+Tab 5 times consecutively", "Screen locks, auto-submits exam", "PASS"),
        ("TC11", "Background Autosave check", "Type descriptive text for 30s", "Draft saves to database as background API", "PASS"),
        ("TC12", "Local Storage backup verification", "Type text, trigger browser crash", "Answers restored from local cache", "PASS"),
        ("TC13", "MCQ Auto Grading correct selection", "Choose correct option D", "Points awarded match allocated marks", "PASS"),
        ("TC14", "MCQ Auto Grading incorrect selection", "Choose incorrect option A", "Awards 0.0 points, correct option logged", "PASS"),
        ("TC15", "Descriptive Grading verbatim match", "Type exact model answer text", "Similarity is 100%, awards full marks", "PASS"),
        ("TC16", "Descriptive Grading semantic similarity", "Type answer using synonyms", "Similarity calculated, scales marks fairly", "PASS"),
        ("TC17", "Descriptive Grading blank field", "Submit empty answer", "Grader awards 0.0 with blank feedback", "PASS"),
        ("TC18", "Live Proctor Monitor active list", "Open admin monitor page", "Displays students taking exams, updates 5s", "PASS"),
        ("TC19", "Live Proctor Warning Alert send", "Click warning, write alert", "Student screen displays warning modal", "PASS"),
        ("TC20", "Supervisor warning beep sound", "Trigger proctor warning modal", "Synthesized alert sound plays on client", "PASS"),
        ("TC21", "Manual submission confirm dialog", "Click submit button", "Popup asks for confirmation", "PASS"),
        ("TC22", "Timer countdown display", "Load active exam page", "Timer starts at set minutes, updates 1s", "PASS"),
        ("TC23", "Timer expiration submit", "Let timer reach 00:00", "Exam submits automatically, locks page", "PASS"),
        ("TC24", "Database transaction safety", "Submit concurrent tests", "MySQL InnoDB writes all scores without error", "PASS"),
        ("TC25", "Password hashing check", "Register new student user", "Password stored as Bcrypt string in database", "PASS")
    ]
    for tc in test_cases_data:
        add_table_row(tbl_test, tc)

    add_subsection_heading(doc, "Detailed Test Scenarios and Results Analysis")
    
    add_paragraph(doc, "1. Authentication and Access Control: We validated the login portals for students and administrators. If a student tries to bypass the login page by writing dashboard URLs directly, PHP redirects them to index.php. Passwords are saved as bcrypt strings in MySQL, protecting candidate profiles. Session tokens are unique and destroyed upon logout, preventing session replay exploits.")
    add_paragraph(doc, "2. Front-End Proctoring Suppression: Test scenarios checked context-menu blocks. Right-clicks inside the container are intercepted and blocked by JS, displaying a warnings toast. Keyboard hooks (Ctrl+C, Ctrl+V, Ctrl+X, Ctrl+A) are intercepted at the body level. Drop events on descriptive textareas are also blocked, preventing text dragging.")
    add_paragraph(doc, "3. Tab Focus Blur Integrity: Tested alt-tab and minimizing events. A blur event increments the count and logs it in the database. Toast warnings are displayed on each blur event. Upon reaching five tab blurs, the exam is immediately locked and submitted. Heartbeat pings update student logs on the proctor dashboard.")
    add_paragraph(doc, "4. Python Grader Calculations: We validated the AI grader on descriptive answers. In TC15, a verbatim match scored 100%. In TC16, an answer written in different words received a scaled score (e.g. 85%) based on keyword overlaps and similarity coefficients. MCQ grading was verified as accurate and instant.")
    add_paragraph(doc, "5. Proctor Warning Broadcasts: The supervisor's warning alerts were tested. When an admin writes a message, it is saved in the database. The student's heartbeat detects this, pauses the exam, and opens a warning modal. The browser plays a synthesized warning sound using the Web Audio API.")

    # ==================== CHAPTER 9: SCREEN SHOTS ====================
    print("Writing Chapter 9 (Screen Shots)...")
    add_chapter_heading(doc, "9 Screen Shots")
    
    add_section_heading(doc, "9.1 Admin Panel Screen Shots")
    add_paragraph(doc, "Because visual layouts are defined in code, the admin dashboard elements are explained in detail below:")
    add_paragraph(doc, "1. Admin Dashboard Home: Displays KPI metric cards in a grid: Total Active Exams, Total Students, Average Score, and Total Violation Warnings. Colors use a professional navy theme with gold accents. A central table lists defined exams, displaying titles, subject codes, duration times, status badges, and action buttons. Admin users can navigate to proctoring panels or exam editor grids easily.")
    add_paragraph(doc, "2. Question Management Interface: Displays a clean form to define questions. Instructors select MCQ or Descriptive type. MCQs show fields for Option A, B, C, D, and correct keys. Descriptive questions show a large textarea for model answers. A list below displays configured questions, showing points and options. Editing questions is disabled once an exam status changes to active or completed.")
    add_paragraph(doc, "3. Real-Time Proctor Monitor Console: The monitoring page displays active students in a responsive card grid. Each card displays candidate names, enrollment numbers, active timers, tab switch counts, and copy-paste attempts. An action bar includes a 'Warn Student' button, opening a prompt to send messages. The page updates automatically every 5 seconds using AJAX.")
    add_paragraph(doc, "4. Detailed Evaluation Sheet: Displays completed student exams. Admins can view individual MCQ choices and descriptive answers. The page shows similarity percentages calculated by the Python grader alongside model answers, allowing instructors to verify automated grades and override scores if necessary.")

    add_section_heading(doc, "9.2 User Panel Screen Shots")
    add_paragraph(doc, "The student dashboard and examination interfaces are configured as follows:")
    add_paragraph(doc, "1. Student Dashboard Home: Displays a clean panel showing student details (name, enrollment number, department). An active exam list displays exam cards, showing subject names, total questions, durations, and green 'Start Assessment' buttons for active exams. Completed exams show a summary scorecard link.")
    add_paragraph(doc, "2. Proctored Exam Interface: The examination page opens in a full-screen layout. The top navbar displays student profiles and connection status indicators. The main section displays one question card. The sidebar navigation palette displays numbered buttons for all questions, indicating answered (green), unanswered (gray), active (blue), and flagged (orange) states.")
    add_paragraph(doc, "3. Proctor Warning Overlay: A modal popup that blocks user interactions when a supervisor warning is received. The window displays warning messages in red text (e.g. '🚨 PROCTOR WARNING: Focus on your screen'). The student must click 'Dismiss & Acknowledge' to return to the exam, logging their acknowledgment in the database.")
    add_paragraph(doc, "4. Student Result Sheet: Shows completed exam summaries, displaying total marks obtained, percentage scores, and individual question feedback. Correct options are displayed, along with similarity scores and feedback generated by the Python engine, providing transparency in subjective grading.")

    # ==================== CHAPTER 10: CONCLUSION ====================
    print("Writing Chapter 10 (Conclusion)...")
    add_chapter_heading(doc, "10 Conclusion")
    add_paragraph(doc, "The Online Examination System with Semantic Evaluation and Real-Time Proctoring successfully establishes a secure, automated platform for remote student assessments at Swaminarayan University. The system addresses the limitations of standard exam tools. The browser lockdown scripts (blocking copy-paste, right-clicks, and tracking tab focus) successfully deter common web-search cheating methods, ensuring exam integrity. The client-side protections are lightweight and run directly inside modern browsers without requiring plugins.")
    add_paragraph(doc, "Additionally, the integration of a Python-based grading script demonstrates that descriptive answers can be graded semantically. Blending sequence matching with keyword overlap creates a reliable evaluation model that aligns with instructor rubrics without requiring verbatim matches. The live admin monitoring dashboard and real-time database autosaves ensure high operational reliability. The system provides a secure, efficient alternative to traditional examinations, reducing administrative overhead.")
    add_paragraph(doc, "The application is scalable and maintains data integrity under concurrent submissions. By automating subjective evaluations and proctoring, universities can conduct examinations remotely with confidence. The platform represents a significant improvement in e-learning tools, showing how deterministic browser controls and natural language processing can be combined to support secure academic assessments.")

    # ==================== CHAPTER 11: FUTURE ENHANCEMENT ====================
    print("Writing Chapter 11 (Future Enhancement)...")
    add_chapter_heading(doc, "11 Future enhancement")
    add_paragraph(doc, "While operational, the platform could be enhanced in several areas to improve security and assessment capabilities:")
    add_paragraph(doc, "1. Visual AI Proctoring: Integrating webcam access to analyze student face positions. The system could use face mesh models in the browser to detect when a student looks away from the screen or leaves their seat, logging these as violations. This would prevent students from referencing books or secondary devices during exams.")
    add_paragraph(doc, "2. Audio Analysis: Accessing candidate microphones during exam sessions. The server could analyze background audio to detect speech activity, flagging potential verbal assistance in the student's room.")
    add_paragraph(doc, "3. Local LLM Integration: Upgrading the Python grader to use local Transformer models or the Gemini API. This would allow the engine to evaluate semantic meaning, logical structure, and code syntax, providing more detailed feedback and grading complex programming questions.")
    add_paragraph(doc, "4. Operating System Lockdown: Developing desktop client applications (e.g., using Electron) to lock down the operating system. This would prevent students from opening secondary applications, virtual machines, or screencasting tools during exams, providing a complete lockdown environment.")
    add_paragraph(doc, "5. Learning Management System (LMS) Integration: Exporting exam data using standard protocols like LTI. This would allow the system to integrate with university platforms like Moodle, Blackboard, or Canvas, synchronizing student registrations and grades automatically.")

    # ==================== CHAPTER 12: BIBLIOGRAPHY ====================
    print("Writing Chapter 12 (Bibliography)...")
    add_chapter_heading(doc, "12 Bibliography")
    add_paragraph(doc, "The following literature references support the research, design, and implementation of this project:")
    
    # 22 Detailed Bibliography References
    refs = [
        "1. Flanagan, D. (2020). JavaScript: The Definitive Guide (7th ed.). O'Reilly Media.",
        "2. Nixon, R. (2021). Learning PHP, MySQL & JavaScript: With jQuery, CSS & HTML5 (6th ed.). O'Reilly Media.",
        "3. Bird, S., Klein, E., & Loper, E. (2019). Natural Language Processing with Python. O'Reilly Media.",
        "4. Sommerville, I. (2018). Software Engineering (10th ed.). Pearson.",
        "5. W3C Document Object Model (DOM) Focus and Page Visibility Specifications. https://www.w3.org/TR/page-visibility/",
        "6. Jurafsky, D., & Martin, J. H. (2023). Speech and Language Processing (3rd ed. draft). Prentice Hall.",
        "7. Pressman, R. S. (2019). Software Engineering: A Practitioner's Approach (9th ed.). McGraw-Hill.",
        "8. Lerdorf, R., Tatroe, K., & MacIntyre, P. (2020). Programming PHP (4th ed.). O'Reilly Media.",
        "9. McKinney, W. (2022). Python for Data Analysis (3rd ed.). O'Reilly Media.",
        "10. DuBois, P. (2018). MySQL (5th ed.). Addison-Wesley Professional.",
        "11. Haverbeke, M. (2021). Eloquent JavaScript: A Modern Introduction to Programming (3rd ed.). No Starch Press.",
        "12. Resig, J., & Bibeault, B. (2020). Secrets of the JavaScript Ninja (2nd ed.). Manning Publications.",
        "13. Fowler, M. (2019). Refactoring: Improving the Design of Existing Code (2nd ed.). Addison-Wesley Professional.",
        "14. Bass, L., Clements, P., & Kazman, R. (2021). Software Architecture in Practice (4th ed.). Addison-Wesley Professional.",
        "15. Thomas, D., & Hunt, A. (2020). The Pragmatic Programmer (20th Anniversary Edition). Addison-Wesley Professional.",
        "16. Kurose, J. F., & Ross, K. W. (2021). Computer Networking: A Top-Down Approach (8th ed.). Pearson.",
        "17. Stallings, W. (2020). Cryptography and Network Security: Principles and Practice (8th ed.). Pearson.",
        "18. Gamma, E., Helm, R., Johnson, R., & Vlissides, J. (2015). Design Patterns: Elements of Reusable Object-Oriented Software. Addison-Wesley Professional.",
        "19. Martin, R. C. (2018). Clean Architecture: A Craftsman's Guide to Software Structure and Design. Prentice Hall.",
        "20. W3C Web Audio API Specification. https://www.w3.org/TR/webaudio/",
        "21. OWASP Top Ten Web Application Security Risks (2021). https://owasp.org/www-project-top-ten/",
        "22. SQLite and Python documentation on SequenceMatcher strings. https://docs.python.org/3/library/difflib.html"
    ]
    for r in refs:
        add_bullet_point(doc, r)

    # Save document
    output_filename = "Online_Examination_System_Thesis.docx"
    print(f"Saving document as {output_filename}...")
    doc.save(output_filename)
    print("Thesis generated successfully!")

if __name__ == '__main__':
    build_thesis_docx()
